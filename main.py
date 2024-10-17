import os
import PyPDF2
import docx
import asyncio
import time
import shutil
import datetime as dt
import streamlit as st
from langchain_core.documents.base import Document
from langchain_community.vectorstores import FAISS
from langchain_pinecone import PineconeEmbeddings
from langchain_groq import ChatGroq
from langchain_core.output_parsers import StrOutputParser
from langchain.prompts import PromptTemplate
from pydrive.drive import GoogleDrive
from pydrive.auth import GoogleAuth

st.title("AI Recruiter")

job = st.text_input("Enter job title*", placeholder="e.g. Graphics Designer")
experience = st.text_input("Enter minimum experience*", placeholder="e.g. 2")
education = st.text_input("Enter minimum education level*", placeholder="e.g. Bachelors")
st.sidebar.header("OPTIONS")
folder = st.sidebar.text_input("Folder Link")
update_button = st.sidebar.button("Update Vector Store")
k = st.sidebar.number_input(label="Number of Resumes to return", min_value=1, max_value=50, step=1, value=10)
t_filter = st.sidebar.selectbox(label="Time Range", options=["1 Hour", "3 Hours", "6 Hours", "12 Hours", "24 Hours", "3 Days", "7 Days", "30 Days", "365 Days", "All"][::-1], placeholder="Select time range", )
keywords = st.text_input("Enter special keywords to search for (optional)", placeholder="e.g. Photoshop, Canva, Figma")

button = st.button("Run")

if t_filter == "1 Hour":
    td = dt.timedelta(seconds=3600)
elif t_filter == "3 Hours":
    td = dt.timedelta(seconds=3*3600)
elif t_filter == "6 Hours":
    td = dt.timedelta(seconds=6*3600)
elif t_filter == "12 Hours":
    td = dt.timedelta(seconds=12*3600)
elif t_filter == "24 Hours":
    td = dt.timedelta(seconds=24*3600)
elif t_filter == "3 Days":
    td = dt.timedelta(seconds=3*24*3600)
elif t_filter == "7 Days":
    td = dt.timedelta(seconds=7*24*3600)
elif t_filter == "30 Days":
    td = dt.timedelta(seconds=30*24*3600)
elif t_filter == "365 Days":
    td = dt.timedelta(seconds=365*24*3600)
else:
    td = None


def update_store():

    gauth = GoogleAuth()
    gauth.LocalWebserverAuth()
    drive = GoogleDrive(gauth)

    os.environ["PINECONE_API_KEY"] = "f03c1dd7-c926-4ea7-851c-bd2723d4fd68"
    folder_id = folder.split("/")[-1]

    embeddings = PineconeEmbeddings(model="multilingual-e5-large")
    if os.path.exists("CV-store"):
        vs = FAISS.load_local("CV-store", embeddings=embeddings, allow_dangerous_deserialization=True)

        vs_files = []

        for doc in vs.docstore.__dict__.values():
            for d in doc.values():
                if d.metadata["folder_id"] == folder_id:
                    vs_files.append(d.metadata["link"])

        no_vs = False
    else:
        no_vs = True

    drive_files = drive.ListFile({'q': f"'{folder_id}' in parents and trashed=false"}).GetList()

    titles = []
    links = []

    os.mkdir("temp")
    for i, f in enumerate(drive_files):
        try:
            file = drive.CreateFile({'id': f["id"]})
            if f["title"].split(".")[-1] in ("pdf", "docx"):
                if not no_vs:
                    if f["title"] not in vs_files:
                        file.GetContentFile("temp/" + f["title"])
                        titles.append("temp/" + f["title"])
                        links.append(f["webContentLink"])
                else:
                    file.GetContentFile("temp/" + f["title"])
                    titles.append("temp/" + f["title"])
                    links.append(f["webContentLink"])
        except Exception:
            pass

        if i % 50 == 0:
            time.sleep(3)

    resumes = []
    new_links = []

    for i, file in enumerate(titles):
        with open(file, "rb") as f:
            error = False
            ftype = file.split(".")[-1]
            try:
                if ftype == "pdf":
                    reader = PyPDF2.PdfReader(f)
                elif ftype == "docx":
                    reader = docx.Document(f)
            except Exception:
                error = True

            if not error and ftype == "pdf":
                resume = ""
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    resume = resume + page.extract_text()
                resumes.append(resume)
                new_links.append(links[i])

            elif not error and ftype == "docx":
                resume = ""
                for para in reader.paragraphs:
                    resume = resume + para.text
                resumes.append(resume)
                new_links.append(links[i])

    shutil.rmtree("temp")

    docs = []

    for i in range(len(resumes)):
        docs.append(Document(page_content=resumes[i],
                             metadata=dict(folder_id=folder_id, timestamp=dt.datetime.now(), link=new_links[i])))

    if len(docs) != 0:
        if no_vs:
            vs = FAISS.from_documents(docs, embeddings)
        else:
            vs.add_documents(docs)

        vs.save_local("CV-store")
    st.sidebar.write(f"Vector store updated!, {len(docs)} new files are added.")


def run():

    start = time.perf_counter()

    os.environ["GROQ_API_KEY"] = "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"

    if not update_button:
        os.environ["PINECONE_API_KEY"] = "f03c1dd7-c926-4ea7-851c-bd2723d4fd68"
        embeddings = PineconeEmbeddings(model="multilingual-e5-large")

        vs = FAISS.load_local("CV-store", embeddings=embeddings, allow_dangerous_deserialization=True)

    if td is None:
        search = vs.similarity_search("A {job} with at least {experience} years of experience with a minimum qualification of "
            "{education}. Keywords: {keywords}", k=k)
    else:
        now = dt.datetime.now()
        search = vs.similarity_search("A {job} with at least {experience} years of experience with a minimum qualification of "
                                      "{education}. Keywords: {keywords}", k=k, filter=lambda x: now - x["timestamp"] < td)
    shortlisted_files = [file.metadata["link"] for file in search]
    shortlisted_timestamps = [file.metadata["timestamp"].strftime("%B %d, %Y") for file in search]
    shortlisted_resumes = [file.page_content for file in search]

    parser = StrOutputParser()

    llm = ChatGroq(
        model="llama-3.1-70b-versatile",
        temperature=0,
        max_tokens=None,
        timeout=None,
        max_retries=2,
    )

    prompt = PromptTemplate.from_template("This is a resume of a candidate {resume}. Extract skills, education,"
                                          "experience of the candidate and phone number and format them in the form of a "
                                          "single markdown table. The table MUST have ONLY 1 row and EXACTLY 6 columns named: "
                                          "Name, City, Skills,"
                                          "Highest level of education, years of experience and phone number. Summarize "
                                          "the name, city, skills,"
                                          "education, experience and phone number in these 5 columns only. Do not output"
                                          "anything else except the table like note or any important point. Just output "
                                          "the table")

    chain = prompt | llm | parser

    if len(shortlisted_resumes) != 0:
        summary = []
        for i in shortlisted_resumes:
            summary.append(chain.invoke({"resume": i}))


        new_summary = []
        for j, r in enumerate(summary):
            new_summary.append(r.split("| --- | --- | --- | --- | --- | --- |")[0].strip() + " Date Added | Link |\n| --- | --- | --- | --- | --- | --- | --- | --- |\n" + r.split("| --- | --- | --- | --- | --- | --- |")[1].strip() + f"{shortlisted_timestamps[j]} | {shortlisted_files[j]} |")

        table = new_summary[0].split("| --- | --- | --- | --- | --- | --- | --- | --- |")[0].strip() + "\n| --- | --- | --- | --- | --- | --- | --- | --- |\n"
        for i in new_summary:
            table += i.split("| --- | --- | --- | --- | --- | --- | --- | --- |")[1].strip() + "\n"

        st.write(table)
    else:
        st.write("No Candidates Found!")

    end = time.perf_counter() - start
    st.write(f"Executed in: {round(end, 2)}s")


async def main():
    run()

async def update():
    update_store()

if update_button:
    if folder:
        asyncio.run(update())
    else:
        st.sidebar.write("Enter link of a folder first!")

if button:
    if job and education and experience:
        asyncio.run(main())
    else:
        st.write("Please enter all the required fields first")